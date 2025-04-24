#!/usr/bin/env python3
"""
Create a test DOCX file for testing the RAG system
"""
from docx import Document

def create_test_docx():
    """Create a sample DOCX file with test content"""
    doc = Document()
    
    doc.add_heading('Sample DOCX Document for Testing', 0)
    
    doc.add_paragraph('This is a sample DOCX document for testing the RAG system.')
    
    doc.add_heading('Section 1: Introduction', 1)
    doc.add_paragraph('This document contains sample content to test the DOCX loading functionality in our RAG CLI tool. '
                     'The tool should be able to ingest this document, split it into chunks, and use it for answering questions.')
    
    doc.add_heading('Section 2: Important Information', 1)
    doc.add_paragraph('The capital of Thailand is Bangkok. It is known for its vibrant street life and ornate shrines. '
                     'The Chao Phraya River flows through the city.')
    
    doc.add_heading('Section 3: More Facts', 1)
    doc.add_paragraph('Thailand has a tropical climate with three main seasons: hot, rainy, and cool. '
                     'The official language is Thai and the currency is the Thai Baht.')
    
    doc.add_heading('Section 4: Conclusion', 1)
    doc.add_paragraph('This document should now be ready for testing our RAG system with DOCX file support.')
    
    # Save the document
    doc.save('data/test_document.docx')
    print("Created test_document.docx in the data directory")

if __name__ == "__main__":
    create_test_docx()
